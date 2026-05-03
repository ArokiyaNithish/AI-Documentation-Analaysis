"""
services/extraction.py — Text extraction from PDF, DOCX, and Images
"""
import io
from pathlib import Path
from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to correct extractor based on file type."""
    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "image": _extract_image,
    }
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"No extractor for file type: {file_type}")
    text = extractor(file_bytes)
    if not text or not text.strip():
        raise ValueError("No text could be extracted from the document.")
    logger.info(f"Extracted {len(text)} characters from {file_type}")
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(layout=True)
                if page_text:
                    text_parts.append(page_text.strip())
                logger.debug(f"PDF page {i+1}: {len(page_text or '')} chars")
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"PDF extraction error: {e}")


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"DOCX extraction error: {e}")


def _extract_image(file_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        from config import settings

        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        image = Image.open(io.BytesIO(file_bytes))
        # Preprocess: convert to grayscale for better OCR accuracy
        image = image.convert("L")
        text = pytesseract.image_to_string(image, config="--psm 6")
        return text.strip()
    except ImportError:
        raise RuntimeError("pytesseract or Pillow not installed.")
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise RuntimeError(f"OCR extraction error: {e}")
